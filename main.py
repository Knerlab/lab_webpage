from flask import Flask, render_template
from flask import redirect, url_for
from docx import Document
import docx
import os
from pdb import set_trace as st
import re
from datetime import datetime
# from wordcloud import WordCloud
import stylecloud
import matplotlib.pyplot as plt

######################################################################################################################################################
'''
functions for extracting content for publications
'''

def extract_year(s):
    # Extract all four-digit numbers from the string
    numbers = re.findall(r'\((\d{4})\)', s)
    
    # Filter out numbers that aren't in the typical range for years
    current_year = datetime.now().year
    years = [int(num) for num in numbers if 1900 <= int(num) <= current_year]
    
    # Return the last found year, as it's more likely to be the publication year
    return years[-1] if years else None


def docx_to_html(filepath):
    document = Document(filepath)
    content = []
    years_list = []

    for para in document.paragraphs:
        cur_year = extract_year(para.text)

        if cur_year not in years_list and cur_year != None:
            years_list.append(cur_year)
            content.append('<div class="years"><p>'+f'{cur_year}'+'</p></div>')

        if para.text != '':
            content.append('<div class="pubs"><p>'+para.text.replace(f' ({cur_year})','').replace(f'[]','')+'</p></div>')
    # st()
        
        # entries.append(para.text)
        
    return " ".join(content)



######################################################################################################################################################
'''
Creating wordcloud using stylecloud
'''

# Function to extract italicized words from the DOCX file
def extract_italic_words_from_docx(filepath):
    document = Document(filepath)
    italic_words = []

    for para in document.paragraphs:
        for run in para.runs:
            if run.italic and run.text.strip():  # Checking if the run is in italic and not just whitespace
                italic_words.append(run.text)

    return ' '.join(italic_words)


if not os.path.exists("staticFiles/files/style_cloud.png"):
    # Extract italic text from the .docx file
    docx_filename = 'staticFiles/files/Publications_website.docx'
    text = extract_italic_words_from_docx(docx_filename).replace('using', '')

    # Generate the word cloud using stylecloud
    stylecloud.gen_stylecloud(text=text,
                            icon_name="fas fa-microscope",
                            palette="colorbrewer.diverging.Spectral_11",
                            background_color='black',
                            gradient="horizontal",
                            max_words=100,
                            output_name="staticFiles/files/style_cloud.png")



######################################################################################################################################################

'''
Flask server
'''
# WSGI Application
# Provide template folder name
# The default folder name should be "templates" else need to mention custom folder name
app = Flask(__name__, template_folder='templateFiles', static_folder='staticFiles')


@app.route('/')
def index():
    return redirect(url_for('home'))

@app.route('/home')
def home():
    return render_template('index.html')

@app.route('/publications')
def publications():
    file_path = os.path.join(app.static_folder, 'files/Publications_website.docx')
    content = docx_to_html(file_path)
    # st()
    return render_template('publications.html', content=content) 

@app.route('/lab-members')
def labMembers():
    return render_template('labMembers.html') 

@app.route('/research')
def research():
    return render_template('research.html') 

@app.route('/teaching-lectures')
def teaching_lectures():
    return render_template('teaching_lectures.html') 

if __name__=='__main__':
    app.run(debug = True, port=8001)