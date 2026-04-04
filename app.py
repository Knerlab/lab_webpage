from flask import Flask, render_template
from flask import redirect, url_for, request, jsonify
from docx import Document
import docx
import os
import csv
import glob
import hashlib
import threading
import ipaddress
from pdb import set_trace as st
import re
from datetime import datetime
from collections import Counter

######################################################################################################################################################
'''
functions for extracting content for publications
'''

def extract_year(s):
    # Extract all four-digit numbers from the string
    numbers = re.findall(r'\((\d{4})\)', s)
    
    # Filter out numbers that aren't in the typical range for years
    current_year = datetime.now().year
    years = [int(num) for num in numbers if 1900 <= int(num) <= current_year]
    
    # Return the last found year, as it's more likely to be the publication year
    return years[-1] if years else None


def docx_to_html(filepath):
    document = Document(filepath)
    content = []
    years_list = []

    for para in document.paragraphs:
        cur_year = extract_year(para.text)

        if cur_year not in years_list and cur_year != None:
            years_list.append(cur_year)
            # Use a specific class 'years' so we can target it in CSS easily
            content.append('<div class="years"><h2>'+f'{cur_year}'+'</h2></div>')

        if para.text != '':
            # The publication text
            content.append('<div class="pubs"><p>'+para.text.replace(f' ({cur_year})','').replace(f'[]','')+'</p></div>')
    # st()
        
        # entries.append(para.text)
        
    return " ".join(content)



######################################################################################################################################################
'''
Creating wordcloud using stylecloud
'''

# Function to extract italicized words from the DOCX file
def extract_italic_words_from_docx(filepath):
    document = Document(filepath)
    italic_words = []

    for para in document.paragraphs:
        for run in para.runs:
            if run.italic and run.text.strip():  # Checking if the run is in italic and not just whitespace
                italic_words.append(run.text)

    return ' '.join(italic_words)


def ensure_style_cloud_exists():
    output_path = os.path.join('staticFiles', 'files', 'style_cloud.png')
    if os.path.exists(output_path):
        return

    docx_filename = os.path.join('staticFiles', 'files', 'Publications_website.docx')
    if not os.path.exists(docx_filename):
        print(f"Warning: {docx_filename} not found. Skipping wordcloud generation.")
        return

    # Lazy import to reduce baseline memory on small servers.
    import stylecloud
    from PIL import ImageDraw

    # Pillow 10+ compatibility for stylecloud internals.
    if not hasattr(ImageDraw.ImageDraw, 'textsize'):
        def textsize(self, text, font=None, *args, **kwargs):
            if font is None:
                font = self.getfont()
            return font.getmask(text).size
        ImageDraw.ImageDraw.textsize = textsize

    text = extract_italic_words_from_docx(docx_filename).replace('using', '')
    stylecloud.gen_stylecloud(
        text=text,
        icon_name="fas fa-microscope",
        palette="colorbrewer.diverging.Spectral_11",
        background_color="rgba(0,0,0,0)",
        gradient="horizontal",
        max_words=100,
        output_name=output_path
    )



######################################################################################################################################################

'''
Flask server
'''
# WSGI Application
# Provide template folder name
# The default folder name should be "templates" else need to mention custom folder name
app = Flask(__name__, template_folder='templateFiles', static_folder='staticFiles')
JOURNALCLUB_DRAW_DIR = os.path.join(app.static_folder, 'journalclub_draw')
JOURNALCLUB_RECORDS_FILE = os.path.join(JOURNALCLUB_DRAW_DIR, 'journalclub_records.csv')
os.makedirs(JOURNALCLUB_DRAW_DIR, exist_ok=True)

ANALYTICS_DIR = os.path.join(os.path.dirname(__file__), 'analytics_data')
ANALYTICS_MAX_ROWS = 15000
ANALYTICS_MAX_RECENT_ROWS = 100
ANALYTICS_ADMIN_TOKEN = os.environ.get('ANALYTICS_ADMIN_TOKEN', '').strip()
ANALYTICS_GEOIP_DB_PATH = os.environ.get(
    'ANALYTICS_GEOIP_DB_PATH',
    os.path.join(ANALYTICS_DIR, 'GeoLite2-City.mmdb')
)
os.makedirs(ANALYTICS_DIR, exist_ok=True)
_analytics_lock = threading.Lock()
_geoip_reader = None
_geoip_checked = False

# ---- Bot / attack filtering ----
# Known bot, crawler, and scanner user-agent keywords (case-insensitive)
_BOT_UA_RE = re.compile(
    r'(bot|crawl|spider|slurp|scanner|python-requests|python-urllib|'
    r'curl|wget|go-http-client|java/|ruby|perl/|'
    r'masscan|nikto|nmap|sqlmap|zgrab|nuclei|httpx|dirsearch|gobuster|'
    r'semrushbot|ahrefsbot|mj12bot|dotbot|baiduspider|yandexbot|'
    r'seznambot|petalbot|bytespider|claudebot|gptbot|chatgpt-user|'
    r'anthropic-ai|dataforseo|paloaltonetworks|nessus|openvas)',
    re.IGNORECASE
)

# Paths typical of automated scanners and vulnerability probes
_ATTACK_PATH_RE = re.compile(
    r'(/\.env|/\.git|/\.ds_store|/wp-|/wordpress|/phpmyadmin|/admin\.php|'
    r'/xmlrpc\.php|/cgi-bin|/shell|/cmd|/config\.|/setup\.|'
    r'/install\.|/backup|/dump|/db\.|/sql\.|'
    r'/v2/|/v1/|/api/v[0-9]|'
    r'/login\.action|/server-status|/server$|/console|/actuator|'
    r'/about$|/info$|/health$|/metrics$|/manager|/solr|/jenkins|'
    r'\.(php|asp|aspx|jsp|cgi|sh|pl|py|rb|cfm)(\?.*)?$)',
    re.IGNORECASE
)


def get_client_ip():
    xff = (request.headers.get('X-Forwarded-For') or '').strip()
    if xff:
        return xff.split(',')[0].strip()

    x_real_ip = (request.headers.get('X-Real-IP') or '').strip()
    if x_real_ip:
        return x_real_ip

    return (request.remote_addr or '').strip()


def mask_ip(ip):
    if not ip:
        return 'unknown'

    if ':' in ip:
        parts = ip.split(':')
        if len(parts) >= 4:
            return ':'.join(parts[:4]) + ':*:*:*:*'
        return ip + ':*'

    parts = ip.split('.')
    if len(parts) == 4:
        return '.'.join(parts[:3]) + '.*'
    return ip


def get_country_city_from_headers():
    country = (
        request.headers.get('CF-IPCountry')
        or request.headers.get('CloudFront-Viewer-Country')
        or request.headers.get('X-Country-Code')
        or request.headers.get('X-Country')
        or 'Unknown'
    ).strip()

    city = (
        request.headers.get('CF-IPCity')
        or request.headers.get('X-City')
        or 'Unknown'
    ).strip()

    return country if country else 'Unknown', city if city else 'Unknown'


def get_geoip_reader():
    global _geoip_reader, _geoip_checked
    if _geoip_checked:
        return _geoip_reader

    _geoip_checked = True
    if not os.path.exists(ANALYTICS_GEOIP_DB_PATH):
        return None

    try:
        import geoip2.database
        _geoip_reader = geoip2.database.Reader(ANALYTICS_GEOIP_DB_PATH)
    except Exception as exc:
        print(f"Warning: GeoIP reader init failed: {exc}")
        _geoip_reader = None

    return _geoip_reader


def get_country_city_from_local_geoip(ip):
    if not ip:
        return 'Unknown', 'Unknown'

    # Ignore local/private addresses.
    try:
        ip_obj = ipaddress.ip_address(ip)
        if ip_obj.is_private or ip_obj.is_loopback or ip_obj.is_reserved:
            return 'Unknown', 'Unknown'
    except ValueError:
        return 'Unknown', 'Unknown'

    reader = get_geoip_reader()
    if reader is None:
        return 'Unknown', 'Unknown'

    try:
        match = reader.city(ip)
        country = (
            (match.country.iso_code or '').strip()
            or (match.country.name or '').strip()
            or 'Unknown'
        )
        city = (match.city.name or '').strip() or 'Unknown'
        return country, city
    except Exception:
        return 'Unknown', 'Unknown'


def build_visitor_fingerprint(ip, user_agent):
    raw = f'{ip}|{user_agent}'
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()[:16]


def analytics_csv_path(now=None):
    now = now or datetime.now()
    return os.path.join(ANALYTICS_DIR, f'visits_{now.strftime("%Y%m")}.csv')


def should_track_request():
    if request.method != 'GET':
        return False

    path = request.path or ''
    if path.startswith('/staticFiles/'):
        return False
    if path.startswith('/api/'):
        return False
    if path.startswith('/admin/analytics'):
        return False

    # Drop known attack/probe paths
    if _ATTACK_PATH_RE.search(path):
        return False

    # Drop empty or known bot/scanner user agents
    ua = (request.headers.get('User-Agent') or '').strip()
    if not ua or _BOT_UA_RE.search(ua):
        return False

    return True


def track_visit(response):
    if not should_track_request():
        return

    # Drop 4xx / 5xx — bot probes almost always land on non-existent paths
    if response.status_code >= 400:
        return

    now = datetime.now()
    path = analytics_csv_path(now)

    ip = get_client_ip()
    user_agent = (request.headers.get('User-Agent') or '').strip()
    referer = (request.headers.get('Referer') or '').strip()
    country, city = get_country_city_from_headers()
    if country == 'Unknown' or city == 'Unknown':
        local_country, local_city = get_country_city_from_local_geoip(ip)
        if country == 'Unknown' and local_country != 'Unknown':
            country = local_country
        if city == 'Unknown' and local_city != 'Unknown':
            city = local_city

    # Skip recording visits with no country info (direct IP access, unresolved geo)
    if country == 'Unknown':
        return

    row = [
        now.strftime('%Y-%m-%d %H:%M:%S'),
        request.path or '',
        request.method,
        str(response.status_code),
        mask_ip(ip),
        build_visitor_fingerprint(ip, user_agent),
        country,
        city,
        referer[:240],
        user_agent[:240]
    ]

    with _analytics_lock:
        exists = os.path.exists(path)
        with open(path, 'a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if not exists:
                writer.writerow([
                    'timestamp',
                    'path',
                    'method',
                    'status',
                    'ip_masked',
                    'visitor_id',
                    'country',
                    'city',
                    'referer',
                    'user_agent'
                ])
            writer.writerow(row)


def load_analytics_rows(limit=ANALYTICS_MAX_ROWS):
    files = sorted(glob.glob(os.path.join(ANALYTICS_DIR, 'visits_*.csv')), reverse=True)
    rows = []

    for file_path in files:
        try:
            with open(file_path, 'r', newline='', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                file_rows = list(reader)
                rows.extend(file_rows[::-1])
        except Exception:
            continue

        if len(rows) >= limit:
            break

    return rows[:limit]


@app.after_request
def set_response_headers(response):
    if request.path.startswith('/staticFiles/'):
        response.headers.setdefault('Cache-Control', 'public, max-age=2592000, immutable')
    elif request.path.startswith('/api/'):
        response.headers.setdefault('Cache-Control', 'no-store')

    try:
        track_visit(response)
    except Exception as exc:
        print(f"Warning: analytics tracking skipped due to error: {exc}")

    return response


@app.route('/')
def index():
    return redirect(url_for('home'))

@app.route('/home')
def home():
    return render_template('index.html')

@app.route('/publications')
def publications():
    file_path = os.path.join(app.static_folder, 'files/Publications_website.docx')
    if os.path.exists(file_path):
        content = docx_to_html(file_path)
    else:
        content = "<p>Publications file not found.</p>"

    try:
        ensure_style_cloud_exists()
    except Exception as exc:
        print(f"Warning: wordcloud generation skipped due to error: {exc}")

    # st()
    return render_template('publications.html', content=content) 

@app.route('/lab-members')
def labMembers():
    return render_template('labMembers.html') 

@app.route('/research')
def research():
    return render_template('research.html') 

@app.route('/teaching-lectures')
def teaching_lectures():
    return render_template('teaching_lectures.html') 

@app.route('/journalclub')
def journalclub():
    return render_template('journalclub.html')


@app.route('/privacy')
def privacy():
    return render_template('privacy.html')


@app.route('/admin/analytics')
def analytics_dashboard():
    if ANALYTICS_ADMIN_TOKEN:
        provided = (request.args.get('token') or '').strip()
        if provided != ANALYTICS_ADMIN_TOKEN:
            return jsonify({'error': 'Unauthorized'}), 401

    rows = load_analytics_rows(limit=ANALYTICS_MAX_ROWS)
    unique_visitors = len({r.get('visitor_id', '') for r in rows if r.get('visitor_id')})

    seen_sessions = set()
    country_counter = Counter()
    city_counter = Counter()
    path_counter = Counter()

    for row in rows:
        vid = row.get('visitor_id', '')
        date = (row.get('timestamp') or '')[:10]
        session_key = f"{vid}|{date}"

        country = (row.get('country') or '').strip()
        city = (row.get('city') or '').strip()
        path = (row.get('path') or '/').strip() or '/'

        path_counter[path] += 1

        if session_key not in seen_sessions:
            seen_sessions.add(session_key)
            if country and country.upper() != 'UNKNOWN':
                country_counter[country] += 1
            if city and city.upper() != 'UNKNOWN':
                city_counter[city] += 1

    total_visits = len(seen_sessions)

    known_rows = [r for r in rows if (r.get('country') or '').strip().upper() != 'UNKNOWN']
    recent_rows = known_rows[:ANALYTICS_MAX_RECENT_ROWS]

    return render_template(
        'analytics.html',
        total_visits=total_visits,
        unique_visitors=unique_visitors,
        country_counts_all=country_counter.most_common(),
        top_countries=country_counter.most_common(10),
        top_cities=city_counter.most_common(10),
        top_paths=path_counter.most_common(10),
        recent_rows=recent_rows,
        token_enabled=bool(ANALYTICS_ADMIN_TOKEN)
    )

@app.route('/api/journalclub/record', methods=['POST'])
def record_journalclub_result():
    payload = request.get_json(silent=True) or {}

    semester = str(payload.get('semester', '')).strip().lower()
    year = str(payload.get('year', '')).strip()
    order = payload.get('order', [])

    if semester not in {'spring', 'winter'}:
        return jsonify({'error': 'Invalid semester.'}), 400

    if not year.isdigit():
        return jsonify({'error': 'Invalid year.'}), 400

    if not isinstance(order, list) or not order:
        return jsonify({'error': 'Invalid order list.'}), 400

    cleaned_order = [str(name).strip() for name in order if str(name).strip()]
    if not cleaned_order:
        return jsonify({'error': 'Order list is empty.'}), 400

    saved_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    csv_exists = os.path.exists(JOURNALCLUB_RECORDS_FILE)

    with open(JOURNALCLUB_RECORDS_FILE, 'a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        if not csv_exists:
            writer.writerow(['saved_at', 'semester', 'year', 'order_count', 'order'])
        writer.writerow([
            saved_at,
            semester,
            year,
            len(cleaned_order),
            ' | '.join(cleaned_order)
        ])

    return jsonify({
        'success': True,
        'saved_at': saved_at,
        'file': 'staticFiles/journalclub_draw/journalclub_records.csv'
    })

@app.route('/api/journalclub/records', methods=['GET'])
def get_journalclub_records():
    if not os.path.exists(JOURNALCLUB_RECORDS_FILE):
        return jsonify({'records': []})

    records = []
    with open(JOURNALCLUB_RECORDS_FILE, 'r', newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            order_raw = (row.get('order') or '').strip()
            order_list = [name.strip() for name in order_raw.split('|') if name.strip()]
            records.append({
                'recordedAt': row.get('saved_at', ''),
                'semester': (row.get('semester') or '').strip(),
                'year': (row.get('year') or '').strip(),
                'order': order_list
            })

    return jsonify({'records': records})

if __name__=='__main__':
    app.run(debug = True, port=8000)
